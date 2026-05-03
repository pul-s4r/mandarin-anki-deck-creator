from __future__ import annotations

import io
from pathlib import Path

from docx import Document


def _docx_document_to_text(document: object) -> str:
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append("\t".join(cells))
    return "\n".join(blocks)


def extract_text_from_docx_bytes(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    return _docx_document_to_text(document)


def extract_text_from_docx(path: Path) -> str:
    return _docx_document_to_text(Document(path))
